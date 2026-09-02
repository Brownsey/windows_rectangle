"""Build the Deep Research DOCX artifact from report-source.md."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "report-source.md"
OUTPUT = ROOT / "research" / "Windows-Rectangle-UX-Research.docx"
LINK = re.compile(r"\[([^]]+)]\((https?://[^)]+)\)")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_rich_text(paragraph, value: str) -> None:
    cursor = 0
    for match in LINK.finditer(value):
        paragraph.add_run(value[cursor : match.start()])
        add_hyperlink(paragraph, match.group(1), match.group(2))
        cursor = match.end()
    paragraph.add_run(value[cursor:])


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in (
        ("Title", 28, "17365D"),
        ("Heading 1", 18, "17365D"),
        ("Heading 2", 13, "2563EB"),
        ("Heading 3", 11, "17365D"),
    ):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    if "Research Meta" not in document.styles:
        meta = document.styles.add_style("Research Meta", WD_STYLE_TYPE.PARAGRAPH)
        meta.font.name = "Aptos"
        meta.font.size = Pt(9)
        meta.font.color.rgb = RGBColor(80, 91, 105)

    header = section.header.paragraphs[0]
    header.text = "WINDOWS RECTANGLE  /  PRODUCT RESEARCH"
    header.style = document.styles["Research Meta"]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("29 AUGUST 2026  •  CONFIDENTIAL WORKING REPORT")
    footer.style = document.styles["Research Meta"]


def add_table(document: Document, lines: list[str]) -> None:
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]
    rows = [row for index, row in enumerate(rows) if index != 1]
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Light Shading Accent 1"
    for index, cell in enumerate(rows[0]):
        table.rows[0].cells[index].text = cell
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, cell in enumerate(row):
            cells[index].text = cell


def build() -> Path:
    document = Document()
    style_document(document)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    first_heading = True
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("| "):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, table_lines)
            continue
        if line.startswith("# "):
            if first_heading:
                paragraph = document.add_paragraph(style="Title")
                add_rich_text(paragraph, line[2:])
                first_heading = False
            else:
                document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_rich_text(paragraph, line[2:])
        elif re.match(r"^\d+\. ", line):
            paragraph = document.add_paragraph(style="List Number")
            add_rich_text(paragraph, re.sub(r"^\d+\. ", "", line))
        elif line.startswith(("Audience:", "Date:", "Scope:")):
            paragraph = document.add_paragraph(style="Research Meta")
            add_rich_text(paragraph, line.removesuffix("  "))
        else:
            paragraph = document.add_paragraph()
            add_rich_text(paragraph, line.replace("**", "").replace("`", ""))
        index += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
